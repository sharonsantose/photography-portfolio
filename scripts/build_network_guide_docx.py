#!/usr/bin/env python3
"""
Generates How_sharonsantos_me_Works.docx on Sharon's Desktop.
Comprehensive architectural & network explanation of sharonsantos.me.
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Georgia'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(16, 185, 129) # Emerald Green
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Georgia'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 41, 59) # Dark Slate
    return p

def add_body_p(doc, text, bold_prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(15, 23, 42)
    r_body = p.add_run(text)
    r_body.font.name = 'Calibri'
    r_body.font.size = Pt(11)
    r_body.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_bullet_p(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(15, 23, 42)
    r_body = p.add_run(text)
    r_body.font.name = 'Calibri'
    r_body.font.size = Pt(11)
    r_body.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F0FDF4") # Light Emerald Tint
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)

    # Left border emerald green
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="10B981"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r_t = p.add_run(title)
    r_t.font.name = 'Calibri'
    r_t.font.size = Pt(11)
    r_t.font.bold = True
    r_t.font.color.rgb = RGBColor(6, 95, 70)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    r_b = p2.add_run(text)
    r_b.font.name = 'Calibri'
    r_b.font.size = Pt(10.5)
    r_b.font.color.rgb = RGBColor(15, 118, 110)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def main():
    doc = docx.Document()

    # Page Margins (1 inch all around)
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    # Document Header Title Block
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    run_t = title_p.add_run("🌐 How sharonsantos.me Works")
    run_t.font.name = 'Georgia'
    run_t.font.size = Pt(26)
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(16, 185, 129)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(16)
    run_s = sub_p.add_run("A Complete Guide to Web Networking, Cloud Automation & System Architecture")
    run_s.font.name = 'Calibri'
    run_s.font.size = Pt(12)
    run_s.font.italic = True
    run_s.font.color.rgb = RGBColor(100, 116, 139)

    meta_table = doc.add_table(rows=1, cols=3)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c0 = meta_table.cell(0, 0)
    c1 = meta_table.cell(0, 1)
    c2 = meta_table.cell(0, 2)
    set_cell_background(c0, "F8FAFC")
    set_cell_background(c1, "F8FAFC")
    set_cell_background(c2, "F8FAFC")
    set_cell_margins(c0, 100, 100, 140, 140)
    set_cell_margins(c1, 100, 100, 140, 140)
    set_cell_margins(c2, 100, 100, 140, 140)

    p_c0 = c0.paragraphs[0]
    p_c0.add_run("AUTHOR: ").bold = True
    p_c0.add_run("Sharon Santos")
    p_c0.runs[0].font.size = Pt(9.5)
    p_c0.runs[1].font.size = Pt(9.5)

    p_c1 = c1.paragraphs[0]
    p_c1.add_run("PLATFORM: ").bold = True
    p_c1.add_run("sharonsantos.me")
    p_c1.runs[0].font.size = Pt(9.5)
    p_c1.runs[1].font.size = Pt(9.5)

    p_c2 = c2.paragraphs[0]
    p_c2.add_run("DATE: ").bold = True
    p_c2.add_run("August 2026")
    p_c2.runs[0].font.size = Pt(9.5)
    p_c2.runs[1].font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # ── SECTION 1: OVERVIEW ──────────────────────────────────────────────────
    add_heading_1(doc, "1. System Overview & Portfolio Structure")
    add_body_p(doc, "sharonsantos.me is a modern, high-performance personal web platform designed to present your creative portfolio (photography, screenwriting treatments, letterboxd cinema reviews) alongside a private control center for your athletic telemetry (Garmin Connect metrics, Strava workouts, and 2:15 Half Marathon training plan).")
    
    add_heading_2(doc, "Core Site Architecture & Page Breakdown:")
    add_bullet_p(doc, "The primary public homepage showcasing your 35mm film photography, essay treatments ('Notes on Light and Motion'), Letterboxd cinema archive, and Strava GPS workouts.", "• sharonsantos.me (index.html): ")
    add_bullet_p(doc, "Your central athletic control center displaying live Garmin health telemetry (Steps, Sleep Score, Resting HR, Training Readiness), Strava workout cards, Google Calendar events, and 2:15 Half Marathon training matrix. Protected by passcode authentication.", "• sharonsantos.me/fit (/fit): ")
    add_bullet_p(doc, "Your daily habits, task management, outdoor weather forecast, and golden hour photography timing calculator. Protected by passcode authentication.", "• sharonsantos.me/todo (/todo): ")
    add_bullet_p(doc, "Legacy private portal configured with an instant client-side rewrite to automatically route traffic cleanly to /fit.", "• sharonsantos.me/private (/private): ")

    # ── SECTION 2: NETWORKING 101 ───────────────────────────────────────────
    add_heading_1(doc, "2. Networking 101 — How Websites Work on the Internet")
    add_body_p(doc, "To understand how sharonsantos.me reaches your phone or laptop from anywhere in the world, it helps to understand the fundamental building blocks of computer networking.")

    add_heading_2(doc, "A. Domain Name System (DNS) — The Internet's Phonebook")
    add_body_p(doc, "Computers on the Internet do not communicate using names like 'sharonsantos.me'. Instead, every connected computer has a unique numerical address called an IP Address (e.g. 185.199.108.153).")
    add_body_p(doc, "When you type 'sharonsantos.me' into your browser, a process called DNS Resolution takes place in a split second:")
    add_bullet_p(doc, "Your web browser asks a DNS recursive resolver (provided by your ISP or Cloudflare 1.1.1.1) 'What is the IP address for sharonsantos.me?'", "1. Request: ")
    add_bullet_p(doc, "The DNS server looks up the domain records (A Records and CNAME Records) configured for your custom domain.", "2. Lookup: ")
    add_bullet_p(doc, "The DNS server returns the IP address pointing to GitHub Pages edge servers (185.199.108.153).", "3. Response: ")
    add_bullet_p(doc, "Your browser establishes a connection directly to that IP address.", "4. Connect: ")

    add_heading_2(doc, "B. The Client-Server Model & HTTP/HTTPS")
    add_body_p(doc, "Once your browser knows the IP address, it acts as a Client and sends an HTTP Request over the network to the Server:")
    add_bullet_p(doc, "An HTTP GET request is sent asking the web server to send back the files needed to render the page (e.g. GET /fit/index.html).", "• HTTP GET Request: ")
    add_bullet_p(doc, "The web server reads the request, retrieves the HTML file, and returns an HTTP 200 OK status code along with the file contents.", "• HTTP 200 OK Response: ")
    add_bullet_p(doc, "HTTPS encrypts all data sent between your browser and the server using SSL/TLS (Secure Sockets Layer). This ensures that passcodes, network packets, and personal telemetry cannot be intercepted by eavesdroppers on public Wi-Fi.", "• HTTPS Encryption: ")

    add_heading_2(doc, "C. Content Delivery Networks (CDNs) & GitHub Pages")
    add_body_p(doc, "sharonsantos.me is hosted on GitHub Pages, which utilizes a global Content Delivery Network (CDN). When a visitor opens your website in San Francisco, Tokyo, or London, GitHub Pages automatically serves the web files from the geographically closest CDN edge data center, resulting in instantaneous sub-100ms page load times.")

    add_callout(doc, "💡 Key Network Takeaway", "When you type sharonsantos.me into your browser, DNS converts the name into an IP address, HTTPS establishes a secure encrypted tunnel, and GitHub's global CDN delivers your site's HTML, CSS, and JS files directly to your device.")

    # ── SECTION 3: FRONTEND ARCHITECTURE & SECURITY ─────────────────────────
    add_heading_1(doc, "3. Frontend Architecture & Security Engineering")
    add_body_p(doc, "The user interface of sharonsantos.me is built entirely using lightweight, vanilla web technologies for zero bloat, high responsiveness, and long-term durability.")

    add_heading_2(doc, "Core Technology Stack:")
    add_bullet_p(doc, "Provides clean, semantic layout structure, accessibility meta tags, and structured DOM elements.", "• HTML5: ")
    add_bullet_p(doc, "Custom high-contrast dark theme (#080808 background), glassmorphism frosted borders, responsive CSS Grid layout, and web fonts (EB Garamond, Inter, Fraunces, JetBrains Mono).", "• Vanilla CSS3: ")
    add_bullet_p(doc, "Client-side logic handling dynamic DOM rendering, real-time cache-busting fetches, workout calculations, and security.", "• Modern JavaScript (ES6+): ")

    add_heading_2(doc, "Passcode Security & Mandatory Refresh Re-Locking:")
    add_body_p(doc, "To protect your personal health data and habits without running an expensive private authentication server, sharonsantos.me uses browser-grade cryptographic SHA-256 hashing:")
    add_bullet_p(doc, "When you type passcode '0427', JavaScript calculates the SHA-256 hash using Web Crypto API (crypto.subtle.digest('SHA-256')).", "1. SHA-256 Hashing: ")
    add_bullet_p(doc, "If the hash matches e4782cfc2b471cd4e24686f692188416d8a313cccb62679dc08c348447c2507b, the content container is unhidden.", "2. Verification: ")
    add_bullet_p(doc, "To guarantee that you are the ONLY person who can view your pages, all unlock session keys (sessionStorage and localStorage) are immediately purged on DOMContentLoaded. Every page refresh forces the passcode screen (authScreen) to show up automatically.", "3. Mandatory Re-Lock on Refresh: ")

    # ── SECTION 4: 100% CLOUD AUTOMATION PIPELINE ────────────────────────────
    add_heading_1(doc, "4. 100% Cloud Automation & Data Pipelines")
    add_body_p(doc, "One of the most powerful features of sharonsantos.me is its ability to update itself automatically without requiring your Mac computer to be powered on or online.")

    add_heading_2(doc, "The Three Automated Data Pipelines:")
    add_bullet_p(doc, "Connects to Garmin Connect API to pull daily steps, sleep score, resting heart rate, active calories, and 30-day telemetry. Saves structured JSON to private/garmin_data.json.", "1. Garmin Health Sync (fetch_garmin.py): ")
    add_bullet_p(doc, "Connects to Strava API v3 via OAuth 2.0 refresh tokens. Retrieves all runs, rides, swims, hikes (e.g. your 9.27 mi hike), and weight workouts, saving to private/strava_data.json.", "2. Strava GPS Workouts Sync (fetch_strava.py): ")
    add_bullet_p(doc, "Uses Google Calendar OAuth 2.0 authorization to query all your calendars (Workouts, Primary, Family) with singleEvents=true expansion. Saves event titles to fit/gcal_imported_events.json.", "3. Google Calendar Sync (fetch_gcal.py): ")

    add_heading_2(doc, "How GitHub Actions Cloud Workflows Operate:")
    add_body_p(doc, "GitHub Actions provides virtual cloud machines running in GitHub's data centers:")
    add_bullet_p(doc, "Workflows in .github/workflows/ (sync-garmin.yml and sync-strava.yml) trigger automatically on a recurring schedule in the cloud.", "• Cloud Trigger: ")
    add_bullet_p(doc, "A virtual Ubuntu machine spins up in the cloud, installs Python, and executes fetch_garmin.py, fetch_strava.py, and fetch_gcal.py.", "• Automated Fetch: ")
    add_bullet_p(doc, "The cloud machine commits the updated JSON files back to your GitHub repository using github-actions[bot].", "• Auto-Commit & Deploy: ")
    add_bullet_p(doc, "GitHub Pages automatically rebuilds the live site. When you view sharonsantos.me/fit on your phone, real-time cache-busting queries (?t=timestamp) load the updated data immediately.", "• Worldwide Publishing: ")

    # ── SECTION 5: VISUAL PIPELINE MAP ──────────────────────────────────────
    add_heading_1(doc, "5. Complete Data & Network Flow Map")
    add_body_p(doc, "Below is the end-to-end flow showing how a workout on your wrist becomes a live card on your website:")

    table_map = doc.add_table(rows=6, cols=3)
    table_map.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["STAGE", "TECHNOLOGY / NODE", "DESCRIPTION"]
    for i, h in enumerate(headers):
        cell = table_map.cell(0, i)
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, 120, 120, 140, 140)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    stages_data = [
        ("1. Telemetry Capture", "Garmin Watch / Strava App", "You record a run, hike, or gym session on your device."),
        ("2. Cloud APIs", "Garmin / Strava / Google APIs", "Activities upload to official cloud provider databases."),
        ("3. Cloud Runner", "GitHub Actions (sync-*.yml)", "Cloud runners execute Python fetch scripts automatically."),
        ("4. Storage & CDN", "GitHub Repository & Pages CDN", "Updated JSON files commit to main and deploy globally."),
        ("5. Client View", "Sharon's Phone / Laptop Browser", "Opening /fit loads real-time telemetry after entering passcode 0427.")
    ]

    for row_idx, data in enumerate(stages_data, start=1):
        bg_hex = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = table_map.cell(row_idx, col_idx)
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, 100, 100, 120, 120)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            r.font.name = 'Calibri'
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(30, 41, 59)
            if col_idx == 0:
                r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    add_callout(doc, "✨ Summary", "Your website is a self-sustaining cloud application. It fetches live data from Garmin, Strava, and Google Calendar 100% autonomously in the cloud, protects your privacy with forced-refresh passcode security, and delivers instant, beautiful telemetry directly to your browser.")

    # Output file destination
    desktop_dir = os.path.expanduser("~/Desktop")
    out_file = os.path.join(desktop_dir, "How_sharonsantos_me_Works.docx")
    doc.save(out_file)
    print(f"✅ Successfully created document at: {out_file}")

if __name__ == "__main__":
    main()
